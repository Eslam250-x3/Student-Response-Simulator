#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
تحويل ملف DOCX إلى TXT
"""

from pathlib import Path
from docx import Document

SCRIPT_DIR = Path(__file__).resolve().parent
DOCX_PATH = SCRIPT_DIR / "الفصل الثالث.docx"
OUTPUT_TXT = SCRIPT_DIR / "الفصل الثالث.txt"


def main():
    if not DOCX_PATH.exists():
        print(f"خطأ: الملف غير موجود: {DOCX_PATH}")
        return 1

    doc = Document(DOCX_PATH)
    lines = []

    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            lines.append(txt)

    # استخراج النص من الجداول إن وجدت
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                lines.append("\t".join(row_text))

    output = "\n".join(lines)
    OUTPUT_TXT.write_text(output, encoding="utf-8")
    print(f"تم الحفظ في: {OUTPUT_TXT}")
    print(f"عدد الأسطر: {len(lines)}")
    return 0


if __name__ == "__main__":
    exit(main())
