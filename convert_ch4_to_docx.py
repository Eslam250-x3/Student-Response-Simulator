import os
import re
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_rtl(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)

def set_font(run, font_name='Simplified Arabic', font_size=14, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    r = run._element
    rFonts = r.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        r.insert(0, rFonts)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

def format_table(table):
    table.style = 'Table Grid'
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for cell in hdr_cells:
        tcPr = cell._element.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'D9D9D9')
        tcPr.append(shading)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_rtl(p)
            for run in p.runs:
                set_font(run, bold=True)

    for row in table.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_rtl(p)
                for run in p.runs:
                    set_font(run, bold=False)

def create_thesis_document(input_txt_file, output_docx_file):
    doc = Document()

    with open(input_txt_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    table_data = []
    in_table = False

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line == "---":
            continue

        if '|' in line:
            in_table = True
            if '---' in line:
                continue # Skip markdown table separator
            row_data = [cell.strip() for cell in line.split('|') if cell.strip()]
            if row_data:
                table_data.append(row_data)
            continue

        if in_table and '|' not in line:
            if table_data:
                cols = len(table_data[0])
                table = doc.add_table(rows=len(table_data), cols=cols)
                table.style = 'Table Grid'
                for i, row in enumerate(table_data):
                    for j, text in enumerate(row):
                        if j < cols:
                            cell = table.cell(i, j)
                            cell.text = text
                format_table(table)
            table_data = []
            in_table = False

        # Clean markdown characters
        text_clean = line
        if text_clean.startswith('> '):
            text_clean = text_clean.replace('> ', '')
        
        # Replace **bold** with bold run (simplistic approach, actually we'll just remove **)
        text_clean = text_clean.replace('**', '')
        text_clean = text_clean.replace('`', '')

        if line.startswith('# '):
            p = doc.add_paragraph()
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text_clean.replace('# ', ''))
            set_font(run, font_size=16, bold=True)
            p.space_after = Pt(12)

        elif line.startswith('## '):
            p = doc.add_paragraph()
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(text_clean.replace('## ', ''))
            set_font(run, font_size=14, bold=True)
            p.space_before = Pt(12)
            p.space_after = Pt(6)

        elif line.startswith('### '):
            p = doc.add_paragraph()
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(text_clean.replace('### ', ''))
            set_font(run, font_size=14, bold=True)
            p.space_before = Pt(10)
            p.space_after = Pt(4)

        elif line.startswith('#### '):
            p = doc.add_paragraph()
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(text_clean.replace('#### ', ''))
            set_font(run, font_size=14, bold=True)
            p.space_before = Pt(8)
            p.space_after = Pt(4)

        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(text_clean[2:])
            set_font(run, font_size=14, bold=False)

        elif re.match(r'^\d+\.', line):
            p = doc.add_paragraph(style='List Number')
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(re.sub(r'^\d+\.\s*', '', text_clean))
            set_font(run, font_size=14, bold=False)

        else:
            p = doc.add_paragraph()
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Inches(0.5)
            run = p.add_run(text_clean)
            set_font(run, font_size=14, bold=False)
            p.paragraph_format.line_spacing = 1.5

    if in_table and table_data:
        cols = len(table_data[0])
        table = doc.add_table(rows=len(table_data), cols=cols)
        table.style = 'Table Grid'
        for i, row in enumerate(table_data):
            for j, text in enumerate(row):
                if j < cols:
                    cell = table.cell(i, j)
                    cell.text = text
        format_table(table)

    doc.save(output_docx_file)
    print(f'تم بنجاح حفظ الملف: {output_docx_file}')

input_file = "fasl_4/الفصل_الرابع_المكتوب.md"
output_file = "fasl_4/الفصل_الرابع_المكتوب.docx"
create_thesis_document(input_file, output_file)
