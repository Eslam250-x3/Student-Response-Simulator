"""
═══════════════════════════════════════════════════════════════
  generate_assignments.py — توليد واجبات الطلاب بالـ AI
  يقرأ simulation_data.json + gradebook ويولّد محتوى لكل طالب/فريق
═══════════════════════════════════════════════════════════════
"""

from __future__ import annotations

import json
import os
import sys
import hashlib
import random
import re
import argparse
import logging
from datetime import datetime
from pathlib import Path
from typing import Any

from api_adapter import APIAdapter

# ─── Logging Setup ──────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger(__name__)

# ─── Constants ──────────────────────────────────────────────
LEARNING_RATIOS = [0.0, 0.40, 0.70, 0.90, 1.0]  # S-curve M1→M5
BASE_DIR = Path(__file__).parent.resolve()


class AssignmentGenerator:
    """Main generator: loads data, builds prompts, calls AI, saves outputs."""

    def __init__(self, config_path: str | Path, require_api: bool = True):
        with open(config_path, 'r', encoding='utf-8') as f:
            self.config: dict = json.load(f)

        self.api = APIAdapter(self.config) if require_api else None
        self.load_data()
        self.ensure_dirs()

    # ═══════════════════════════════════════════════════════
    #  Data Loading
    # ═══════════════════════════════════════════════════════

    def load_data(self) -> None:
        """Load all required data files."""
        sim_path = BASE_DIR / self.config['input']['simulation_data']
        const_path = BASE_DIR / self.config['input']['constants']
        examples_path = BASE_DIR / 'examples_pool.json'

        with open(sim_path, 'r', encoding='utf-8') as f:
            self.students_data: list[dict] = json.load(f)['students']
        self.student_by_id: dict[str, dict] = {s['id']: s for s in self.students_data}

        with open(const_path, 'r', encoding='utf-8') as f:
            self.constants: dict = json.load(f)
        with open(examples_path, 'r', encoding='utf-8') as f:
            self.examples_pool: dict = json.load(f)

        # Load Gradebook for target quality
        self.gradebook: dict[str, dict] = {}
        gradebook_path = BASE_DIR / self.config['input']['gradebook_csv']
        if gradebook_path.exists():
            import csv
            with open(gradebook_path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    self.gradebook[row['ID']] = {
                        "name": row.get('Name', ''),
                        "group": row.get('Group', ''),
                        "team": row.get('Team', ''),
                        "percentage": float(row['Percentage']) if row.get('Percentage') else 0,
                        "grade": row.get('Grade', ''),
                        "lateness": {
                            f"M{i}": row.get(f'M{i}_Late', 'لا') for i in range(1, 6)
                        },
                        "dates": {
                            f"M{i}": row.get(f'M{i}_Date', '') for i in range(1, 6)
                        },
                    }

        self._build_team_mappings_from_gradebook()

        # Load task descriptions
        self.tasks_info: dict[str, dict] = {}
        for m in ['M1', 'M2', 'M3', 'M4', 'M5']:
            task_path = BASE_DIR / 'tasks' / f'{m}.json'
            with open(task_path, 'r', encoding='utf-8') as f:
                self.tasks_info[m] = json.load(f)

        logger.info(f"Loaded {len(self.students_data)} students, {len(self.gradebook)} gradebook entries")

    def ensure_dirs(self) -> None:
        """Create output directories (relative to script, not cwd)."""
        docx_root = BASE_DIR / 'outputs' / 'docx'
        docx_root.mkdir(parents=True, exist_ok=True)
        for group in ['G1', 'G2', 'G3', 'G4']:
            (docx_root / group).mkdir(parents=True, exist_ok=True)
        (BASE_DIR / 'outputs' / 'json').mkdir(parents=True, exist_ok=True)

    # ═══════════════════════════════════════════════════════
    #  Student Profile Generation
    # ═══════════════════════════════════════════════════════

    @staticmethod
    def get_seed(identifier: str) -> int:
        """Deterministic seed from any string identifier."""
        return int(hashlib.md5(identifier.encode()).hexdigest(), 16) % (10**8)

    @staticmethod
    def _parse_datetime(value: str) -> datetime | None:
        value = (value or "").strip()
        if not value:
            return None
        for fmt in ("%Y-%m-%d %H:%M", "%Y-%m-%d %H:%M:%S"):
            try:
                return datetime.strptime(value, fmt)
            except ValueError:
                continue
        return None

    @staticmethod
    def _team_code(group: str, team_label: str) -> str:
        """
        Build stable team code (e.g. G3_Team_6) from gradebook team label.
        """
        raw = (team_label or "").strip()
        if not raw or raw == "عمل فردي":
            return f"{group}_Team_unknown"

        digits = re.findall(r"\d+", raw)
        if digits:
            return f"{group}_Team_{digits[0]}"

        normalized = re.sub(r"\s+", "_", raw)
        normalized = re.sub(r"[^\w\u0600-\u06FF-]+", "_", normalized)
        normalized = normalized.strip("_")
        return f"{group}_Team_{normalized or 'unknown'}"

    def _build_team_mappings_from_gradebook(self) -> None:
        """
        Build collaborative team members and per-milestone submitter from gradebook.
        Submitter is chosen as the member with earliest submission datetime.
        """
        self.team_members: dict[str, list[dict]] = {}
        self.team_labels: dict[str, str] = {}
        self.team_submitter_by_milestone: dict[tuple[str, str], str] = {}

        for sid, gdata in self.gradebook.items():
            group = gdata.get("group", "")
            if group not in ["G3", "G4"]:
                continue

            team_label = (gdata.get("team") or "").strip() or "unknown"
            team_code = self._team_code(group, team_label)
            self.team_labels[team_code] = team_label

            student = self.student_by_id.get(sid, {
                "id": sid,
                "name": gdata.get("name", sid),
                "group": group,
                "preSkill": 0.5,
                "postSkill": 0.5,
                "preFlowLevel": 0.5,
                "postFlowLevel": 0.5,
            })
            self.team_members.setdefault(team_code, []).append(student)

        milestones = ["M1", "M2", "M3", "M4", "M5"]
        for team_code, members in self.team_members.items():
            member_ids = sorted({m.get("id", "") for m in members if m.get("id")})
            if not member_ids:
                continue

            for milestone in milestones:
                dated: list[tuple[datetime, str]] = []
                for sid in member_ids:
                    dt_value = self._parse_datetime(
                        self.gradebook.get(sid, {}).get("dates", {}).get(milestone, "")
                    )
                    if dt_value:
                        dated.append((dt_value, sid))

                if dated:
                    dated.sort(key=lambda x: x[0])
                    submitter_sid = dated[0][1]
                else:
                    submitter_sid = member_ids[0]

                self.team_submitter_by_milestone[(team_code, milestone)] = submitter_sid

    @staticmethod
    def _is_female(name: str) -> bool:
        """
        Heuristic: Detect gender from Arabic name.
        Since all students in the current cohort are female, this returns True
        if name doesn't belong to a known male set (extend as needed).
        """
        male_hints = ["أحمد", "محمد", "علي", "عبد", "محمود", "إبراهيم", "ياسر", "مصطفى", "خالد"]
        # If it's a female name like "نورهان أحمد", the first name is the key.
        first_name = name.split()[0].strip() if name else ""
        
        # known female names in this cohort
        female_names = [
            "نورهان", "سارة", "مريم", "هدى", "آية", "رنا", "دينا", "نهى", "شيماء", "إسراء", 
            "ريهام", "منى", "سلمى", "لمياء", "أميرة", "حنان", "عبير", "رشا", "علا", "نجلاء",
            "فاطمة", "زينب", "هبة", "ندى", "روان", "جنى", "بسمة", "ملك", "تقى", "لجين"
        ]
        if first_name in female_names:
            return True
        # Check against male first names
        if first_name in male_hints and first_name not in ["نور"]: # 'noor' can be both
            return False
            
        return True # Default to female for this cohort

    def generate_student_profile(self, student_id: str) -> dict:
        """Generate consistent profile (writing style, interest, M3 selection)."""
        seed = self.get_seed(student_id)
        random.seed(seed)
        styles = self.config['realism']['writingStyles']
        interests = self.config['realism'].get('interests', [])
        topics = list(self.examples_pool.keys())
        return {
            "writing_style": random.choice(styles),
            "interest": random.choice(interests) if interests else "عام",
            "m3_selection": random.sample(topics, min(3, len(topics))),
        }

    # ═══════════════════════════════════════════════════════
    #  Prompt Building
    # ═══════════════════════════════════════════════════════

    def build_prompt(
        self,
        student: dict,
        milestone: str,
        profile: dict,
        skill: float,
    ) -> tuple[str, str]:
        """
        Build system_prompt and user_prompt for a specific student+milestone.
        Returns (system_prompt, user_prompt).
        """
        task = self.tasks_info[milestone]
        is_time_pressure = student['group'] in ['G2', 'G4']

        is_female = self._is_female(student.get('name', ''))

        # ─── System Prompt (Persona & Constraints) ───
        if is_female:
            parts = [
                "أنتِ طالبة مصرية في الصف الأول الثانوي (15-16 سنة).",
                "اكتبي بأسلوب مراهقة ذكية تحاول حل واجبها المدرسي بجدية.",
                "استخدمي لغة عربية فصحى بسيطة مناسبة لعمركِ.",
                "تجنبي استخدام مصطلحات معقدة جداً أو احترافية إلا إذا كنتِ تستشهدين بمصدر.",
                "اجعلي نبرتكِ تعكس تساؤلات وميول جيلكِ.",
            ]
        else:
            parts = [
                "أنت طالب مصري في الصف الأول الثانوي (15-16 سنة).",
                "اكتب بأسلوب مراهق ذكي يحاول حل واجبه المدرسي بجدية.",
                "استخدم لغة عربية فصحى بسيطة مناسبة لعمرك.",
                "تجنب استخدام مصطلحات معقدة جداً أو احترافية إلا إذا كنت تستشهد بمصدر.",
                "اجعل نبرتك تعكس تساؤلات وميول جيلك.",
            ]

        # Quality Constraints (Critical Fix Phase 7)
        parts.extend([
            "\n**قواعد هامة جداً للمخرجات (توقفي عن فعل الآتي):**",
            "1. اكتبي الإجابة مباشرة وبالعربية فقط. لا تكتبي أي تخطيط أو تفكير أو مسودات بالإنجليزي أو العربي.",
            "2. لا تستخدمي علامات التنسيق (bold **) أو العناوين (#) أو أي markdown. اكتبي كأنكِ تكتبي في ورقة واجب عادية.",
            "3. لا تضعي النص بين علامات اقتباس.",
            "4. لا تكتبي ملاحظات جانبية أو عد كلمات أو 'Let's draft'. فقط النص النهائي.",
        ])

        # Skill-based adjustment
        if skill > 0.7:
            parts.append("أنتِ طالبة متميزة، كتابتكِ منظمة وعميقة ودقيقة." if is_female else "أنت طالب متميز، كتابتك منظمة وعميقة ودقيقة.")
        elif skill < 0.4:
            parts.append("أنتِ طالبة مستواكِ ضعيف، قد ترتكبين بعض الأخطاء الإملائية البسيطة وتكون جملكِ غير مكتملة أحياناً." if is_female else "أنت طالب مستواك ضعيف، قد ترتكب بعض الأخطاء الإملائية البسيطة وتكون جملك غير مكتملة أحياناً.")
        else:
            parts.append("أنتِ طالبة متوسطة المستوى، شرحكِ واضح وبسيط." if is_female else "أنت طالب متوسط المستوى، شرحك واضح وبسيط.")

        # Writing style & interest
        parts.append(f"أسلوبكِ في الكتابة هو: {profile['writing_style']}." if is_female else f"أسلوبك في الكتابة هو: {profile['writing_style']}.")
        parts.append(f"أنتِ طالبة مهتمة بـ {profile['interest']}، لذا حاولي ربط أفكاركِ بهذا الاهتمام كلما أمكن بشكل طبيعي." if is_female else f"أنت طالب مهتم بـ {profile['interest']}، لذا حاول ربط أفكارك بهذا الاهتمام كلما أمكن بشكل طبيعي.")

        # Flow state
        flow_idx = student.get('flow_idx', 0.5)
        if flow_idx > 0.8:
            parts.append("أنتِ الآن في حالة تدفق ذهني عالية، تشعرين بالتركيز الشديد والاستمتاع والاندماج التام." if is_female else "أنت الآن في حالة تدفق ذهني عالية، تشعر بالتركيز الشديد والاستمتاع والاندماج التام.")
        elif flow_idx < 0.3:
            parts.append("أنتِ تشعرين ببعض التشتت أو الملل، وقد تبدو كتابتكِ أقل حماساً أو تفتقر للتفاصيل." if is_female else "أنت تشعر ببعض التشتت أو الملل، وقد تبدو كتابتك أقل حماساً أو تفتقر للتفاصيل.")

        # Lateness persona
        target = self.gradebook.get(student['id'], {})
        if target:
            is_late = target.get('lateness', {}).get(milestone, 'لا') == 'نعم'
            if is_late:
                parts.append("أنتِ تسلمين المهمة بعد الموعد النهائي، لذا قد تبدو نبرتكِ معتذرة قليلاً أو متوترة بسبب التأخير." if is_female else "أنت تسلم المهمة بعد الموعد النهائي، لذا قد تبدو نبرتك معتذرة قليلاً أو متوترة بسبب التأخير.")

        system_prompt = " ".join(parts)

        # ─── User Prompt (Content) ───
        prompt_lines = []

        # M1 Welcome Context (Based on group)
        if milestone == 'M1':
            welcome_map = {
                "G1": "🚀 مرحباً بك في مستودع المعرفة التنافسي! أنت هنا لتثبت تميزك الفردي وتساهم بأفضل ما لديك في هذا التحدي الفلسفي.",
                "G2": "⚡ مرحباً بك في تحدي حشد المصادر السريع! السرعة والجودة هما مفتاح فوزك في هذا السباق مع الزمن.",
                "G3": "🤝 مرحباً بكم في رحلة التعاون المعرفي! قوتكم في اتحادكم وتبادل خبراتكم لبناء دليل شامل للأخلاق.",
                "G4": "⏰ مرحباً بكم في تحدي التعاون السريع! النجاح يعتمد على تناغم فريقكم وإنجاز المهام الجماعية بدقة وتحت الضغط.",
            }
            prompt_lines.append(welcome_map.get(student['group'], ""))
            prompt_lines.append("")

        # Target quality from gradebook
        if target:
            prompt_lines.append(f"**درجة الجودة المستهدفة لعملك هي: {target.get('percentage', 0):.0f}% (تقدير {target.get('grade', '')})**")
            is_late = target.get('lateness', {}).get(milestone, 'لا') == 'نعم'
            sub_date = target.get('dates', {}).get(milestone, '')
            if is_late:
                prompt_lines.append(f"**تنبيه: أنت تسلم هذا العمل متأخراً (تاريخ التسليم الفعلي: {sub_date})**")
            prompt_lines.append("")

        # Task description & Instructions
        prompt_lines.append(f"الوصف: {task['description']}")
        
        # Add Rubric to prompt (Phase 7)
        prompt_lines.append("\nمعايير التقييم التي يجب مراعاتها:")
        for criterion, points in task.get('rubric', {}).items():
            prompt_lines.append(f"- {criterion}: {points} نقاط")
        prompt_lines.append("")

        # Instructions (shuffled for diversity)
        instructions = task['instructions'].copy()
        random.seed(self.get_seed(student['id'] + milestone))
        random.shuffle(instructions)
        prompt_lines.append("التعليمات التفصيلية:")
        for inst in instructions:
            prompt_lines.append(f"- {inst}")

        # Time pressure
        prompt_lines.append("")
        if is_time_pressure:
            prompt_lines.append("**تنبيه ضيق الوقت:** تذكر أن مجموعتك (G2/G4) لديها مواعيد نهائية صارمة جداً. اكتب بتركيز وسرعة.")
        else:
            prompt_lines.append("**تنبيه الإتقان:** مجموعتك (G1/G3) لديها وقت كافٍ. استثمر الوقت في البحث والتنسيق الجيد.")

        # M2 Specific: pooled examples
        if milestone == 'M2':
            prompt_lines.append("\nاستلهم من الأحداث التالية (صغها بأسلوبك الشخصي كبحث قمت به):")
            for category, examples in self.examples_pool.items():
                if examples:
                    ex = random.choice(examples)
                    prompt_lines.append(f"- {category}: {ex['title']} ({ex['date']}) - {ex['description']}")

        # M3 Specific: topic selection
        if milestone == 'M3':
            prompt_lines.append("\nمهم جداً: ابدأ مباشرة بعنوان **القضية الأولى** بدون أي مقدمة أو بسملة أو تمهيد.")
            prompt_lines.append("لا تكتب أي فقرة افتتاحية قبل القضايا الثلاث.")
            prompt_lines.append(f"\nالتحليل المطلوب للقضايا التالية التي اخترتها: {', '.join(profile['m3_selection'])}.")

        # M5 Specific: compilation logic (Phase 7)
        if milestone == 'M5':
            prompt_lines.append("\nهذا هو عملك النهائي التجميعي. يجب أن يتضمن:")
            prompt_lines.append("- مقدمة (100 كلمة) تعبر عن رحلتك في تعلم الأخلاق البيوطبية.")
            prompt_lines.append("- تجميع ومنسق للمهام M1, M2, M3, M4 بعد مراجعتها.")
            prompt_lines.append("- خاتمة (150 كلمة) تلخص أهم الدروس المستفادة.")

        # Word count (Stronger emphasis Phase 7)
        prompt_lines.append(f"\nتنبيه هام جداً: يجب أن يكون طول النص بين {task['minWords']}-{task['maxWords']} كلمة بالضبط.")
        prompt_lines.append(f"اكتبي نصاً طويلاً وغنياً بالتفاصيل، سيتم رفض الإجابة لو كانت قصيرة.")

        prompt = "\n".join(prompt_lines)
        return system_prompt, prompt

    # ═══════════════════════════════════════════════════════
    #  Output Validation
    # ═══════════════════════════════════════════════════════

    def validate_output(self, content: str, milestone: str) -> tuple[bool, str]:
        """
        Validate generated content against task requirements.
        Returns (is_valid, reason).
        """
        if not self.config['realism'].get('validateOutput', False):
            return True, "Validation disabled"

        task = self.tasks_info[milestone]
        word_count = len(content.split())
        min_w = int(task['minWords'] * 0.7)  # 30% tolerance below
        max_w = int(task['maxWords'] * 1.4)  # 40% tolerance above

        if word_count < min_w:
            return False, f"Too short: {word_count} words (min ~{min_w})"
        if word_count > max_w:
            return False, f"Too long: {word_count} words (max ~{max_w})"

        # M1: should have 4 sections (paragraphs)
        if milestone == 'M1':
            paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            if len(paragraphs) < 3:
                return False, f"M1 should have ~4 sections, found {len(paragraphs)}"

        # M2: should mention 5 topics
        if milestone == 'M2':
            topics_found = sum(1 for topic in self.examples_pool.keys() if topic in content)
            if topics_found < 3:
                return False, f"M2 should cover 5 topics, found {topics_found}"

        # M3: must start directly with issue #1 (no intro/preamble)
        if milestone == 'M3':
            first_non_empty = next((line.strip() for line in content.splitlines() if line.strip()), "")
            normalized = re.sub(r"^[\-\*\d\.\)\(\s]+", "", first_non_empty)
            if "القضية الأولى" not in normalized:
                return False, "M3 should start directly with 'القضية الأولى' (no intro)"

        return True, f"OK ({word_count} words)"

    @staticmethod
    def _strip_m3_intro(content: str) -> str:
        """
        Remove any preface before the first M3 issue heading.
        Keeps text unchanged if no issue heading is found.
        """
        lines = content.splitlines()
        for idx, line in enumerate(lines):
            if "القضية الأولى" in line:
                return "\n".join(lines[idx:]).strip()
        return content.strip()

    @staticmethod
    def _safe_path_name(value: str, fallback: str = "Unknown") -> str:
        """Return a filesystem-safe path component while preserving Arabic letters."""
        text = (value or "").strip()
        text = re.sub(r'[\\/:*?"<>|]+', "_", text)
        text = re.sub(r"\s+", " ", text).strip()
        return text if text else fallback

    def _build_docx_output_path(self, metadata: dict) -> Path:
        """Build nested DOCX path: outputs/docx/<group>/<submitter>/<milestone_date>.docx"""
        group_folder = self._safe_path_name(str(metadata.get('group', 'UNKNOWN')), fallback='UNKNOWN')
        submitter_name = metadata.get('submitter_name') or metadata.get('sname') or metadata.get('sid') or 'Unknown'
        submitter_folder = self._safe_path_name(str(submitter_name), fallback='Unknown')
        milestone = self._safe_path_name(str(metadata.get('milestone', 'M')), fallback='M')
        sub_date = metadata.get('submission_date') or metadata.get('date') or 'NoDate'
        sub_date_safe = self._safe_path_name(str(sub_date), fallback='NoDate')
        filename = f"{milestone}_{sub_date_safe}.docx"
        return BASE_DIR / 'outputs' / 'docx' / group_folder / submitter_folder / filename

    # ═══════════════════════════════════════════════════════
    #  DOCX Generation
    # ═══════════════════════════════════════════════════════

    def save_docx(self, content: str, metadata: dict) -> Path:
        """Save content as a formatted Word document."""
        try:
            from docx import Document
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.warning("python-docx not installed, skipping DOCX generation")
            return None

        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Title
        title = doc.add_heading(level=1)
        title_run = title.add_run(f"{metadata.get('milestone', '')} — {metadata.get('submitter_name', metadata.get('sname', ''))}")
        title_run.font.size = Pt(18)
        title_run.font.name = 'Simplified Arabic'
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Metadata
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        meta_text = f"المجموعة: {metadata.get('group', '')} | "
        meta_text += f"التاريخ: {metadata.get('submission_date', '')} | "
        meta_text += f"المستوى: {metadata.get('skill', 0):.0%}"
        meta_run = meta_para.add_run(meta_text)
        meta_run.font.size = Pt(10)
        meta_run.font.name = 'Simplified Arabic'
        meta_run.font.color.rgb = None  # Default color

        participant_names = metadata.get('participant_names', [])
        if participant_names:
            members_para = doc.add_paragraph()
            members_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            members_text = "الطلاب المشاركون: " + "، ".join(participant_names)
            members_run = members_para.add_run(members_text)
            members_run.font.size = Pt(10)
            members_run.font.name = 'Simplified Arabic'

        doc.add_paragraph()  # spacer

        # Content
        for paragraph_text in content.split('\n'):
            if paragraph_text.strip():
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = para.add_run(paragraph_text.strip())
                run.font.size = Pt(14)
                run.font.name = 'Simplified Arabic'

        # Save
        output_path = self._build_docx_output_path(metadata)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.debug(f"Saved DOCX: {output_path}")
        return output_path

    # ═══════════════════════════════════════════════════════
    #  Token Estimation
    # ═══════════════════════════════════════════════════════

    @staticmethod
    def estimate_max_tokens(task: dict) -> int:
        """Estimate max_output_tokens from task word requirements."""
        max_words = task.get('maxWords', 500)
        # Arabic averages ~1.5 tokens per word, add 30% buffer
        return min(int(max_words * 1.5 * 1.3), 8192)

    # ═══════════════════════════════════════════════════════
    #  Dry Run
    # ═══════════════════════════════════════════════════════

    def dry_run(self, limit: int | None = None, student_filter: str | None = None,
                milestone_filter: str | None = None) -> None:
        """Print generation plan without calling API."""
        jobs = self._build_job_list(limit, student_filter, milestone_filter)

        print("\n" + "═" * 60)
        print("  🧪 DRY RUN — خطة التوليد (بدون استدعاء API)")
        print("═" * 60)
        print(f"\n📊 إجمالي الواجبات: {len(jobs)}")

        # Group summary
        groups = {}
        for job in jobs:
            g = job['group']
            groups[g] = groups.get(g, 0) + 1
        print("\n📋 التوزيع حسب المجموعة:")
        for g, count in sorted(groups.items()):
            print(f"   {g}: {count} واجب")

        # Milestone summary
        milestones = {}
        for job in jobs:
            m = job['milestone']
            milestones[m] = milestones.get(m, 0) + 1
        print("\n📋 التوزيع حسب المهمة:")
        for m, count in sorted(milestones.items()):
            print(f"   {m}: {count} واجب")

        # Sample prompt
        if jobs:
            sample = jobs[0]
            s = sample['student']
            profile = self.generate_student_profile(s['id'])
            m_idx = int(sample['milestone'][1]) - 1
            skill = s['preSkill'] + (s['postSkill'] - s['preSkill']) * LEARNING_RATIOS[m_idx]
            s_task = s.copy()
            s_task['flow_idx'] = s.get('preFlowLevel', 0.5)
            sys_p, p = self.build_prompt(s_task, sample['milestone'], profile, skill)

            print(f"\n{'─' * 60}")
            print(f"📝 عينة Prompt — {sample['job_id']}:")
            print(f"{'─' * 60}")
            print(f"\n🔧 System Prompt ({len(sys_p)} حرف):")
            print(f"   {sys_p[:200]}...")
            print(f"\n📨 User Prompt ({len(p)} حرف):")
            print(f"   {p[:300]}...")
            print(f"\n🎯 Skill: {skill:.2f} | Style: {profile['writing_style']}")
            max_tokens = self.estimate_max_tokens(self.tasks_info[sample['milestone']])
            print(f"🔢 Max tokens: {max_tokens}")

        print("\n" + "═" * 60)
        print("  ✅ Dry run complete — لم يتم استدعاء أي API")
        print("═" * 60 + "\n")

    # ═══════════════════════════════════════════════════════
    #  Job List Builder
    # ═══════════════════════════════════════════════════════

    def _build_job_list(
        self,
        limit: int | None = None,
        student_filter: str | None = None,
        milestone_filter: str | None = None,
    ) -> list[dict]:
        """Build list of all jobs to process, deduplicating teams."""
        dropout_ids = set(self.constants.get('DROPOUT_IDS', []) + self.constants.get('dropoutIds', []))
        seen_team_jobs: set[str] = set()
        jobs: list[dict] = []

        for s in self.students_data:
            sid = s['id']
            group = s['group']

            # Student filter
            if student_filter and sid != student_filter:
                continue

            is_dropout = sid in dropout_ids
            milestones = ['M1', 'M2'] if is_dropout else ['M1', 'M2', 'M3', 'M4', 'M5']

            for m in milestones:
                # Milestone filter
                if milestone_filter and m != milestone_filter:
                    continue

                # Build job_id (team-based for G3/G4)
                if group in ['G3', 'G4']:
                    gb_team = self.gradebook.get(sid, {}).get('team', '')
                    team_id = self._team_code(group, gb_team)
                    team_label = self.team_labels.get(team_id, gb_team or "unknown")
                    job_id = f"{team_id}_{m}"
                    if job_id in seen_team_jobs:
                        continue
                    seen_team_jobs.add(job_id)
                else:
                    team_id = None
                    team_label = None
                    job_id = f"{sid}_{m}"

                jobs.append({
                    "job_id": job_id,
                    "student": s,
                    "milestone": m,
                    "group": group,
                    "team_id": team_id,
                    "team_label": team_label,
                    "is_dropout": is_dropout,
                })

                if limit and len(jobs) >= limit:
                    return jobs

        return jobs

    # ═══════════════════════════════════════════════════════
    #  Main Run
    # ═══════════════════════════════════════════════════════

    def run(
        self,
        limit: int | None = None,
        student_filter: str | None = None,
        milestone_filter: str | None = None,
    ) -> None:
        """Generate assignments by calling AI API."""
        if self.api is None:
            self.api = APIAdapter(self.config)

        # Load progress
        progress_path = BASE_DIR / 'progress.json'
        if progress_path.exists():
            with open(progress_path, 'r', encoding='utf-8') as f:
                progress: dict = json.load(f)
        else:
            progress = {}

        # Build jobs
        jobs = self._build_job_list(None, student_filter, milestone_filter)
        total_jobs = len(jobs)
        pending_all = [j for j in jobs if j['job_id'] not in progress]
        pending_jobs = pending_all[:limit] if limit is not None else pending_all

        logger.info(
            f"Total jobs: {total_jobs} | Already done: {total_jobs - len(pending_all)} | "
            f"Pending: {len(pending_all)} | This run: {len(pending_jobs)}"
        )

        if not pending_jobs:
            logger.info("✅ All jobs already completed!")
            return

        outputs_json: dict = {}
        completed = 0

        for i, job in enumerate(pending_jobs, 1):
            base_student = job['student']
            sid = base_student['id']
            sname = base_student.get('name', 'Unknown')
            m = job['milestone']
            group = job['group']
            team_id = job['team_id']
            team_label = job.get('team_label')
            job_id = job['job_id']

            submitter_sid = sid
            if group in ['G3', 'G4'] and team_id:
                submitter_sid = self.team_submitter_by_milestone.get((team_id, m), sid)

            s = self.student_by_id.get(submitter_sid, base_student)
            sid = s.get('id', submitter_sid)
            sname = self.gradebook.get(sid, {}).get('name') or s.get('name', 'Unknown')

            logger.info(f"[{i}/{len(pending_jobs)}] Generating {job_id}...")

            # Calculate skill & flow
            m_idx = int(m[1]) - 1
            skill = s['preSkill'] + (s['postSkill'] - s['preSkill']) * LEARNING_RATIOS[m_idx]
            flow = s.get('preFlowLevel', 0.5) + (s.get('postFlowLevel', 0.5) - s.get('preFlowLevel', 0.5)) * LEARNING_RATIOS[m_idx]

            s_task = s.copy()
            s_task['flow_idx'] = flow

            profile = self.generate_student_profile(sid)
            sys_p, p = self.build_prompt(s_task, m, profile, skill)

            # Dynamic temperature & seed
            job_seed = self.get_seed(job_id)
            temp = 0.6 + (job_seed % 25) / 100.0  # Range 0.6 to 0.85
            max_tokens = self.estimate_max_tokens(self.tasks_info[m])

            try:
                content = self.api.call_ai(
                    p, sys_p,
                    temperature=temp,
                    seed=job_seed,
                    max_output_tokens=max_tokens,
                )
                if m == 'M3':
                    content = self._strip_m3_intro(content)

                # Validate
                is_valid, reason = self.validate_output(content, m)
                if not is_valid:
                    logger.warning(f"Validation failed for {job_id}: {reason}. Retrying...")
                    retry_prompt = p + f"\n\n**تنبيه: الإجابة السابقة {reason}. أعد الكتابة بالطول المناسب.**"
                    content = self.api.call_ai(
                        retry_prompt, sys_p,
                        temperature=temp + 0.05,
                        seed=job_seed + 1,
                        max_output_tokens=max_tokens,
                    )
                    if m == 'M3':
                        content = self._strip_m3_intro(content)
                    is_valid2, reason2 = self.validate_output(content, m)
                    if not is_valid2:
                        logger.warning(f"Retry validation: {reason2}. Using anyway.")

                # File naming
                sub_date = self.gradebook.get(sid, {}).get('dates', {}).get(m, 'NoDate')
                sub_date_safe = sub_date.replace('/', '-').replace(':', '-').replace(' ', '_')
                file_id = f"{sname}_{m}_{sub_date_safe}"
                if group in ['G3', 'G4']:
                    file_id = f"{team_id}_{m}_{sub_date_safe}"

                participant_names: list[str] = []
                if group in ['G3', 'G4'] and team_id:
                    members = self.team_members.get(team_id, [])
                    participant_names = sorted({
                        member.get('name', 'Unknown').strip() or 'Unknown'
                        for member in members
                    })

                # Metadata
                metadata = {
                    "sid": sid,
                    "sname": sname,
                    "submitter_name": sname,
                    "group": group,
                    "team": team_id,
                    "team_label": team_label,
                    "milestone": m,
                    "date": sub_date,
                    "skill": round(skill, 3),
                    "flow": round(flow, 3),
                    "style": profile['writing_style'],
                    "interest": profile['interest'],
                    "is_late": self.gradebook.get(sid, {}).get('lateness', {}).get(m, 'لا'),
                    "submission_date": sub_date,
                    "word_count": len(content.split()),
                    "temperature": round(temp, 3),
                    "participant_names": participant_names,
                }

                # Store result
                outputs_json[file_id] = {
                    "content": content,
                    "metadata": metadata,
                    "prompts": {
                        "system": sys_p,
                        "user": p,
                    },
                }

                # Save DOCX if enabled
                if self.config['output'].get('docx_enabled', False):
                    self.save_docx(content, metadata)

                # Mark progress
                progress[job_id] = True
                completed += 1
                with open(progress_path, 'w', encoding='utf-8') as f:
                    json.dump(progress, f)

                logger.info(f"  ✅ {job_id} — {len(content.split())} words, skill={skill:.2f}")

            except Exception as e:
                logger.error(f"  ❌ Failed {job_id}: {e}")
                continue

        # Save all results as JSON
        json_path = BASE_DIR / 'outputs' / 'json' / 'assignments.json'
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(outputs_json, f, ensure_ascii=False, indent=2)

        logger.info(f"\n🎉 Generation complete! {completed}/{len(pending_jobs)} succeeded.")
        logger.info(f"   JSON: {json_path}")
        if self.config['output'].get('docx_enabled', False):
            logger.info(f"   DOCX: {BASE_DIR / 'outputs' / 'docx'}")


# ═══════════════════════════════════════════════════════════════
#  CLI
# ═══════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="توليد واجبات الطلاب بالـ AI",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
أمثلة الاستخدام:
  python generate_assignments.py --dry-run           # معاينة الخطة
  python generate_assignments.py --limit 5           # توليد 5 فقط
  python generate_assignments.py --student STD-001   # طالب واحد
  python generate_assignments.py --milestone M3      # مهمة واحدة
  python generate_assignments.py --reset             # حذف progress
        """,
    )
    parser.add_argument('--dry-run', action='store_true', help='عرض خطة التوليد بدون استدعاء API')
    parser.add_argument('--limit', type=int, default=None, help='أقصى عدد واجبات (بدون حد = الكل)')
    parser.add_argument('--student', type=str, default=None, help='توليد لطالب واحد فقط (مثل STD-001)')
    parser.add_argument('--milestone', type=str, default=None, choices=['M1','M2','M3','M4','M5'], help='توليد مهمة واحدة فقط')
    parser.add_argument('--reset', action='store_true', help='حذف progress.json والبدء من الصفر')
    parser.add_argument('--config', type=str, default=None, help='مسار ملف config.json بديل')
    parser.add_argument('--verbose', '-v', action='store_true', help='طباعة تفاصيل إضافية')

    args = parser.parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    if args.reset:
        progress_path = BASE_DIR / 'progress.json'
        if progress_path.exists():
            progress_path.unlink()
            logger.info("🗑️  تم حذف progress.json")
        else:
            logger.info("ℹ️  progress.json غير موجود")
        return

    config_path = args.config or (BASE_DIR / 'config.json')
    gen = AssignmentGenerator(config_path, require_api=not args.dry_run)

    if args.dry_run:
        gen.dry_run(limit=args.limit, student_filter=args.student, milestone_filter=args.milestone)
    else:
        gen.run(limit=args.limit, student_filter=args.student, milestone_filter=args.milestone)


if __name__ == "__main__":
    main()
