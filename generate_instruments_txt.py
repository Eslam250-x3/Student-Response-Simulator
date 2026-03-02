# -*- coding: utf-8 -*-
"""
توليد ملفات الاختبار ومقياس التدفق الذهني في مجلد مقاييس نهائيه
يدعم: txt, json, docx (بتنسيق احترافي)
"""
import json
import os
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = SCRIPT_DIR / "مقاييس نهائيه"
TEST_CONFIG = SCRIPT_DIR / "test_config.json"
FLOW_CONFIG = SCRIPT_DIR / "flow_config.json"

LABELS = ["أ", "ب", "ج", "د"]

# ─── رأس وتعليمات الاختبار ───
TEST_HEADER = """الاختبار مهارات حل المشكلات في الفلسفة البيوطبية

بيانات الطالب:
الاسم: ___________________________________
الفصل: ___________________________________
التاريخ: ___________________________________

تعليمات الاختبار:

عزيزي الطالب/الطالبة،

أولاً: طبيعة الاختبار:
يهدف هذا الاختبار إلى قياس مهاراتك في حل المشكلات من خلال مواقف أخلاقية بيوطبية.
يتكون الاختبار من 30 سؤالاً من نوع الاختيار من متعدد.
كل سؤال عبارة عن موقف حياتي أو دراسي يتبعه أربعة بدائل.

ثانياً: الهدف من الاختبار:
قياس قدرتك على تحديد المشكلات الأخلاقية.
قياس قدرتك على افتراض أسباب محتملة للمشكلات.
قياس قدرتك على اختبار صحة الفروض بطرق منطقية.
قياس قدرتك على الوصول إلى أفضل الحلول للمشكلات الأخلاقية.

ثالثاً: كيفية الإجابة:
اقرأ كل موقف بعناية وتأنٍ.
اختر إجابة واحدة فقط من البدائل الأربعة (أ، ب، ج، د).
ضع علامة (✓) أمام الإجابة التي تعبر عن استجابتك.

رابعاً: الزمن المخصص:
زمن الاختبار: 40 دقيقة.

خامساً: ملاحظات هامة:
لا تترك أي سؤال دون إجابة.
أجب بصدق، فالاختبار يعبر عن وجهة نظرك.
جميع البيانات التي يتم جمعها ستُستخدم بسرية تامة لأغراض البحث العلمي فقط.

مثال توضيحي:
مريض يحتاج لعملية جراحية عاجلة وموصى بها طبيًا، لكنه يرفض إجراءها بسبب خوفه الشديد وغير المبرر من التخدير، مما يضعه في خطر. ما هي المشكلة الأساسية التي يجب التعامل معها؟
أ) التكلفة المالية العالية للعملية الجراحية.
ب) نقص الأطباء المتخصصين في هذا النوع من الجراحات.
ج) ✓ الصراع بين حاجة المريض للعلاج وحقه في الرفض بناءً على خوفه.
د) عدم وجود غرف عمليات كافية في المستشفى حاليًا.
الإجابة المناسبة هي (ج) لأنها تحدد المشكلة الأساسية بدقة.

جدول توزيع الأسئلة على المهارات
م	المهارات	أرقام المفردات	المجموع	النسبة
1	تحديد المشكلة	1-4-6-11-16-21-26	7	23.33%
2	فرض الفروض	2-7-9-12-17-22-27	7	23.33%
3	اختبار صحة الفروض	3-8-13-14-18-23-28	7	23.33%
4	الوصول للحل الصحيح	5-10-15-19-20-24-25-29-30	9	30%
	المجموع			30	100%

أسئلة الاختبار

"""

# ─── رأس وتعليمات المقياس ───
FLOW_HEADER = """مقياس التدفق الذهني

عزيزي الطالب / عزيزتي الطالبة،

بين يديك مجموعة من العبارات التي تصف شعورك أثناء التفاعل وحل المشكلات الأخلاقية في بيئة التعلم الإلكتروني.
يرجى قراءة كل عبارة بدقة، واختيار ما ينطبق عليك تماماً.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📌 دليل الإجابة:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

◉ دائماً: تنطبق عليك العبارة بصفة مستمرة
◉ غالباً: تنطبق عليك العبارة في معظم الأحيان
◉ أحياناً: تنطبق عليك العبارة بعض الوقت
◉ نادراً: تنطبق عليك العبارة في أوقات قليلة جداً
◉ أبداً: لا تنطبق عليك العبارة إطلاقاً

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⏱️ الوقت المتوقع: 10-15 دقيقة
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

العبارات

"""

# ─── دوال مساعدة لـ DOCX ───
FONT_ARABIC = "Simplified Arabic"


def _set_paragraph_rtl(paragraph):
    """تفعيل اتجاه يمين-لشمال للفقرة."""
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        pPr.append(bidi)
    except Exception:
        pass


def _add_arabic_paragraph(doc, text, font_size=14, bold=False):
    """إضافة فقرة عربية منسقة."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_rtl(para)
    run = para.add_run(text)
    run.font.name = FONT_ARABIC
    run.font.size = Pt(font_size)
    run.font.bold = bold
    para.paragraph_format.line_spacing = 1.5
    return para


def generate_test_docx():
    """توليد ملف الاختبار .docx بتنسيق احترافي."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[WARN] python-docx غير مثبت، تخطي توليد DOCX")
        return

    with open(TEST_CONFIG, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # العنوان الرئيسي
    _add_arabic_paragraph(doc, "الاختبار مهارات حل المشكلات في الفلسفة البيوطبية", font_size=18, bold=True)
    doc.add_paragraph()

    # بيانات الطالب
    _add_arabic_paragraph(doc, "بيانات الطالب:", font_size=14, bold=True)
    _add_arabic_paragraph(doc, "الاسم: ___________________________________", font_size=14)
    _add_arabic_paragraph(doc, "الفصل: ___________________________________", font_size=14)
    _add_arabic_paragraph(doc, "التاريخ: ___________________________________", font_size=14)
    doc.add_paragraph()

    # تعليمات الاختبار
    _add_arabic_paragraph(doc, "تعليمات الاختبار:", font_size=16, bold=True)
    _add_arabic_paragraph(doc, "عزيزي الطالب/الطالبة،", font_size=14)
    _add_arabic_paragraph(doc, "أولاً: طبيعة الاختبار: يهدف هذا الاختبار إلى قياس مهاراتك في حل المشكلات من خلال مواقف أخلاقية بيوطبية. يتكون الاختبار من 30 سؤالاً من نوع الاختيار من متعدد. كل سؤال عبارة عن موقف حياتي أو دراسي يتبعه أربعة بدائل.", font_size=14)
    _add_arabic_paragraph(doc, "ثانياً: الهدف من الاختبار: قياس قدرتك على تحديد المشكلات الأخلاقية، افتراض أسباب محتملة، اختبار صحة الفروض، والوصول إلى أفضل الحلول.", font_size=14)
    _add_arabic_paragraph(doc, "ثالثاً: كيفية الإجابة: اقرأ كل موقف بعناية، اختر إجابة واحدة من البدائل الأربعة (أ، ب، ج، د)، وضع علامة (✓) أمام الإجابة.", font_size=14)
    _add_arabic_paragraph(doc, "رابعاً: الزمن المخصص: 40 دقيقة.", font_size=14)
    _add_arabic_paragraph(doc, "خامساً: ملاحظات هامة: لا تترك أي سؤال دون إجابة. أجب بصدق. جميع البيانات سرية لأغراض البحث العلمي فقط.", font_size=14)
    doc.add_paragraph()

    # مثال توضيحي
    _add_arabic_paragraph(doc, "مثال توضيحي:", font_size=14, bold=True)
    _add_arabic_paragraph(doc, "مريض يحتاج لعملية جراحية عاجلة وموصى بها طبيًا، لكنه يرفض إجراءها بسبب خوفه الشديد وغير المبرر من التخدير، مما يضعه في خطر. ما هي المشكلة الأساسية التي يجب التعامل معها؟", font_size=14)
    _add_arabic_paragraph(doc, "أ) التكلفة المالية العالية للعملية الجراحية.", font_size=14)
    _add_arabic_paragraph(doc, "ب) نقص الأطباء المتخصصين في هذا النوع من الجراحات.", font_size=14)
    _add_arabic_paragraph(doc, "ج) ✓ الصراع بين حاجة المريض للعلاج وحقه في الرفض بناءً على خوفه.", font_size=14)
    _add_arabic_paragraph(doc, "د) عدم وجود غرف عمليات كافية في المستشفى حاليًا.", font_size=14)
    _add_arabic_paragraph(doc, "الإجابة المناسبة هي (ج) لأنها تحدد المشكلة الأساسية بدقة.", font_size=14)
    doc.add_paragraph()

    # جدول توزيع المهارات
    _add_arabic_paragraph(doc, "جدول توزيع الأسئلة على المهارات", font_size=14, bold=True)
    table = doc.add_table(rows=6, cols=5)
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    headers = ["م", "المهارات", "أرقام المفردات", "المجموع", "النسبة"]
    rows_data = [
        ("1", "تحديد المشكلة", "1-4-6-11-16-21-26", "7", "23.33%"),
        ("2", "فرض الفروض", "2-7-9-12-17-22-27", "7", "23.33%"),
        ("3", "اختبار صحة الفروض", "3-8-13-14-18-23-28", "7", "23.33%"),
        ("4", "الوصول للحل الصحيح", "5-10-15-19-20-24-25-29-30", "9", "30%"),
        ("", "المجموع", "", "30", "100%"),
    ]
    for c, h in enumerate(headers):
        cell = table.rows[0].cells[c]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in p.runs:
                r.font.name = FONT_ARABIC
                r.font.size = Pt(12)
                r.font.bold = True
    for r_idx, row_data in enumerate(rows_data, start=1):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for r in p.runs:
                    r.font.name = FONT_ARABIC
                    r.font.size = Pt(12)
    doc.add_paragraph()

    # أسئلة الاختبار
    _add_arabic_paragraph(doc, "أسئلة الاختبار", font_size=16, bold=True)
    doc.add_paragraph()

    for q in data["questions"]:
        idx = q["order"]
        text = q["text"]
        choices = q["choices"]
        _add_arabic_paragraph(doc, f"[{idx}] {text}", font_size=14, bold=True)
        for i, choice in enumerate(choices):
            _add_arabic_paragraph(doc, f"    {LABELS[i]}) {choice}", font_size=14)
        doc.add_paragraph()

    # مفتاح الإجابات
    _add_arabic_paragraph(doc, "مفتاح الإجابات:", font_size=14, bold=True)
    key_parts = [f"Q{q['order']}:{q['correctLabel']}" for q in data["questions"]]
    _add_arabic_paragraph(doc, " ".join(key_parts[:10]), font_size=12)
    _add_arabic_paragraph(doc, " ".join(key_parts[10:20]), font_size=12)
    _add_arabic_paragraph(doc, " ".join(key_parts[20:30]), font_size=12)

    out_path = OUTPUT_DIR / "اختبار_المشكلات_الأخلاقية_البيوطبية.docx"
    doc.save(str(out_path))
    print(f"[OK] تم إنشاء: {out_path}")


def generate_flow_docx():
    """توليد ملف المقياس .docx بتنسيق احترافي."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[WARN] python-docx غير مثبت، تخطي توليد DOCX")
        return

    with open(FLOW_CONFIG, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # العنوان
    _add_arabic_paragraph(doc, "مقياس التدفق الذهني", font_size=18, bold=True)
    doc.add_paragraph()

    # الترحيب والتعليمات
    _add_arabic_paragraph(doc, "عزيزي الطالب / عزيزتي الطالبة،", font_size=14)
    _add_arabic_paragraph(doc, "بين يديك مجموعة من العبارات التي تصف شعورك أثناء التفاعل وحل المشكلات الأخلاقية في بيئة التعلم الإلكتروني. يرجى قراءة كل عبارة بدقة، واختيار ما ينطبق عليك تماماً.", font_size=14)
    doc.add_paragraph()

    _add_arabic_paragraph(doc, "دليل الإجابة:", font_size=14, bold=True)
    _add_arabic_paragraph(doc, "دائماً: تنطبق عليك العبارة بصفة مستمرة", font_size=14)
    _add_arabic_paragraph(doc, "غالباً: تنطبق عليك العبارة في معظم الأحيان", font_size=14)
    _add_arabic_paragraph(doc, "أحياناً: تنطبق عليك العبارة بعض الوقت", font_size=14)
    _add_arabic_paragraph(doc, "نادراً: تنطبق عليك العبارة في أوقات قليلة جداً", font_size=14)
    _add_arabic_paragraph(doc, "أبداً: لا تنطبق عليك العبارة إطلاقاً", font_size=14)
    doc.add_paragraph()

    _add_arabic_paragraph(doc, "الوقت المتوقع: 10-15 دقيقة", font_size=14, bold=True)
    doc.add_paragraph()

    # العبارات
    _add_arabic_paragraph(doc, "العبارات", font_size=16, bold=True)
    doc.add_paragraph()

    for item in data["items"]:
        idx = item["id"]
        text = item["text"]
        sign = "[+]" if not item["isNegative"] else "[-]"
        _add_arabic_paragraph(doc, f"{idx}. {sign} {text}", font_size=14)

    out_path = OUTPUT_DIR / "مقياس_التدفق_الذهني.docx"
    doc.save(str(out_path))
    print(f"[OK] تم إنشاء: {out_path}")


def generate_test_txt():
    """توليد ملف الاختبار .txt"""
    with open(TEST_CONFIG, "r", encoding="utf-8") as f:
        data = json.load(f)

    lines = [TEST_HEADER]
    for q in data["questions"]:
        idx = q["order"]
        text = q["text"]
        choices = q["choices"]
        correct = q["correctAnswer"]
        lines.append(f"[{idx}] {text}")
        for i, choice in enumerate(choices):
            mark = " ✓" if i == correct else ""
            lines.append(f"    {LABELS[i]}) {choice}{mark}")
        lines.append("")

    lines.append("\n───────────────────────────────────────────────────────────────────────")
    lines.append("مفتاح الإجابات:")
    lines.append("───────────────────────────────────────────────────────────────────────")
    key_parts = []
    for q in data["questions"]:
        key_parts.append(f"Q{q['order']}:{q['correctLabel']}")
    lines.append(" ".join(key_parts[:10]))
    lines.append(" ".join(key_parts[10:20]))
    lines.append(" ".join(key_parts[20:30]))

    out_path = OUTPUT_DIR / "اختبار_المشكلات_الأخلاقية_البيوطبية.txt"
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"[OK] تم إنشاء: {out_path}")


def generate_test_json():
    """نسخ الاختبار إلى .json"""
    with open(TEST_CONFIG, "r", encoding="utf-8") as f:
        data = json.load(f)
    out_path = OUTPUT_DIR / "اختبار_المشكلات_الأخلاقية_البيوطبية.json"
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"[OK] تم إنشاء: {out_path}")


def generate_flow_txt():
    """توليد ملف المقياس .txt"""
    with open(FLOW_CONFIG, "r", encoding="utf-8") as f:
        data = json.load(f)

    lines = [FLOW_HEADER]
    for item in data["items"]:
        idx = item["id"]
        text = item["text"]
        dim = item["dimension"]
        sign = "[+]" if not item["isNegative"] else "[-]"
        lines.append(f"{idx}. {sign} {text}")

    out_path = OUTPUT_DIR / "مقياس_التدفق_الذهني.txt"
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"[OK] تم إنشاء: {out_path}")


def generate_flow_json():
    """نسخ المقياس إلى .json"""
    with open(FLOW_CONFIG, "r", encoding="utf-8") as f:
        data = json.load(f)
    out_path = OUTPUT_DIR / "مقياس_التدفق_الذهني.json"
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"[OK] تم إنشاء: {out_path}")


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    generate_test_txt()
    generate_test_json()
    generate_flow_txt()
    generate_flow_json()
    generate_test_docx()
    generate_flow_docx()
    print("\n[OK] تم إنشاء جميع الملفات بنجاح في مجلد: مقاييس نهائيه")


if __name__ == "__main__":
    main()
