"""Extract text from docx and PDF research documents."""
import sys, os

# 1. Read DOCX
from docx import Document
doc = Document(r'e:\New folder (6)\الاطار سوبر نهائي.docx')
with open(r'e:\New folder (6)\_docx_out.txt', 'w', encoding='utf-8') as f:
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            f.write(txt + '\n')
print("DOCX done:", len(doc.paragraphs), "paragraphs")

# 2. Read PDF  
from PyPDF2 import PdfReader
reader = PdfReader(r'e:\New folder (6)\اسلام خطة بحث معدل نسخة ما قبل التسليم للمجلس معدل سلوى اخر نسخة بعد التنسيق .pdf')
with open(r'e:\New folder (6)\_pdf_out.txt', 'w', encoding='utf-8') as f:
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            f.write(f'\n--- PAGE {i+1} ---\n')
            f.write(text + '\n')
print("PDF done:", len(reader.pages), "pages")
