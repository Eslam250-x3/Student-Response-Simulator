# -*- coding: utf-8 -*-
"""
تحديث المواعيد النهائية في ملف المجموعة الرابعة (تعاوني بضغط زمني)
من: 4, 9, 15, 18, 21
إلى: 3, 8, 14, 17, 21 (متسق مع generate_tasks_gradebook.py)
"""
import os
from docx import Document

magmo3at_dir = os.path.join(os.path.dirname(__file__), 'magmo3at')
g4_file = os.path.join(magmo3at_dir, 'المجموعة الرابعة_ تعاوني بضغط زمني.docx')

# استبدالات: (قديم, جديد)
replacements = [
    ('يوم 4', 'يوم 3'),
    ('اليوم 4', 'اليوم 3'),
    ('يوم 9', 'يوم 8'),
    ('اليوم 9', 'اليوم 8'),
    ('يوم 15', 'يوم 14'),
    ('اليوم 15', 'اليوم 14'),
    ('يوم 18', 'يوم 17'),
    ('اليوم 18', 'اليوم 17'),
    # تجنب استبدال "يوم 21" - صحيح
    ('أيام 4، 9، 15، 18، 21', 'أيام 3، 8، 14، 17، 21'),
    ('اليوم 1-4:', 'اليوم 1-3:'),
    ('اليوم 5-9:', 'اليوم 4-8:'),
    ('اليوم 10-15:', 'اليوم 9-14:'),
    ('اليوم 16-18:', 'اليوم 15-17:'),
    ('اليوم 19-21:', 'اليوم 18-21:'),
    ('(4 أيام)', '(3 أيام)'),
    ('[4 أيام فقط!]', '[3 أيام فقط!]'),
    ('يوم 4، 11:59', 'يوم 3، 11:59'),
    ('يوم 9، 11:59', 'يوم 8، 11:59'),
    ('يوم 15، 11:59', 'يوم 14، 11:59'),
    ('يوم 18، 11:59', 'يوم 17، 11:59'),
    ('يوم 4، الساعة', 'يوم 3، الساعة'),
    ('يوم 9، الساعة', 'يوم 8، الساعة'),
    ('يوم 15، الساعة', 'يوم 14، الساعة'),
    ('يوم 18، الساعة', 'يوم 17، الساعة'),
    ('يوم 4، 6 مساءً', 'يوم 3، 6 مساءً'),
    ('يوم 4، الساعة 6', 'يوم 3، الساعة 6'),
    ('يوم 4، 11:59 م', 'يوم 3، 11:59 م'),
    ('يوم 9، 11:59 م', 'يوم 8، 11:59 م'),
    ('يوم 15، 11:59 م', 'يوم 14، 11:59 م'),
    ('يوم 18، 11:59 م', 'يوم 17، 11:59 م'),
]

def replace_in_element(element, old, new):
    """Replace text in paragraph or similar element (handles split runs)."""
    if not hasattr(element, 'text') or old not in element.text:
        return 0
    count = element.text.count(old)
    new_text = element.text.replace(old, new)
    element.clear()
    element.add_run(new_text)
    return count

def iter_paragraphs(doc):
    """Yield all paragraphs in document and tables."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def main():
    if not os.path.exists(g4_file):
        print("File not found")
        return
    
    doc = Document(g4_file)
    total_changes = 0
    
    for old, new in replacements:
        for para in iter_paragraphs(doc):
            n = replace_in_element(para, old, new)
            total_changes += n
    
    doc.save(g4_file)
    print("Done. Total replacements:", total_changes)

if __name__ == '__main__':
    main()
